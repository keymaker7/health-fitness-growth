# -*- coding: utf-8 -*-
"""Build Word files Copilot Studio will accept as knowledge."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "agents" / "upload"


def set_font(run, size=11):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, 16 if level == 1 else 13)


def add_para(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run)
    run.bold = bold
    return p


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            for p in table.rows[i].cells[j].paragraphs:
                for run in p.runs:
                    set_font(run, 10)
                    if i == 0:
                        run.bold = True


def md_to_docx(md_text: str, title: str) -> Document:
    doc = Document()
    add_heading(doc, title, 1)
    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    table_buf: list[str] = []

    def flush_table():
        nonlocal table_buf
        rows = []
        for raw in table_buf:
            if re.match(r"^\|?\s*-{3,}", raw.replace("|", " | ")):
                continue
            if set(raw.replace("|", "").replace(" ", "").replace(":", "")) <= {"-"}:
                continue
            cells = [c.strip() for c in raw.strip().strip("|").split("|")]
            if cells:
                rows.append(cells)
        if rows:
            add_table(doc, rows)
        table_buf = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        if table_buf:
            flush_table()
        if line.startswith("# "):
            # already have title; skip duplicate h1
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3)
        elif line.startswith("> "):
            add_para(doc, line[2:].strip())
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            for run in p.runs:
                set_font(run)
        elif line.startswith("```"):
            i += 1
            chunk = []
            while i < len(lines) and not lines[i].startswith("```"):
                chunk.append(lines[i])
                i += 1
            add_para(doc, "\n".join(chunk))
        elif line.strip() == "" or line.strip() == "---":
            pass
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line).strip()
            if text:
                add_para(doc, text)
        i += 1
    if table_buf:
        flush_table()
    return doc


def paps_doc() -> Document:
    events = json.loads((ROOT / "src" / "data" / "paps-events.json").read_text(encoding="utf-8"))
    doc = Document()
    add_heading(doc, "PAPS 종목 안내 (앱 동기화본)", 1)
    add_para(
        doc,
        "교육청 원본 매뉴얼 파일이 없을 때 쓰는 지식입니다. "
        "앱 src/data/paps-events.json과 같습니다. "
        "출처는 학교건강검사규칙 별표 3과 제주특별자치도교육청 학생건강체력평가(PAPS) 운영 매뉴얼(2019)입니다. "
        "문서에 없는 종목은 안내하지 않습니다.",
    )
    add_para(doc, "줄넘기는 이 목록에 없습니다. PAPS 필수 종목이 아니라 수업 연습 도구입니다.", bold=True)

    for ev in events:
        add_heading(doc, f"{ev['name']} ({ev['fitnessFactor']})", 2)
        add_para(doc, f"목적: {ev['purpose']}")
        add_para(doc, f"적용: {ev['applicable']} / 단위: {ev['unit']} / 장소: {ev['place']}")
        add_para(doc, "준비물: " + ", ".join(ev["tools"]))
        add_para(doc, "실시 방법", bold=True)
        for m in ev["method"]:
            p = doc.add_paragraph(m, style="List Number")
            for run in p.runs:
                set_font(run)
        add_para(doc, "자세", bold=True)
        for m in ev["posture"]:
            p = doc.add_paragraph(m, style="List Bullet")
            for run in p.runs:
                set_font(run)
        add_para(doc, "주의", bold=True)
        for m in ev["cautions"]:
            p = doc.add_paragraph(m, style="List Bullet")
            for run in p.runs:
                set_font(run)
        src = " / ".join(s["title"] for s in ev.get("sources", []))
        if src:
            add_para(doc, f"근거: {src}")
    return doc


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    files = [
        ("01-class-knowledge.docx", ROOT / "docs" / "agents" / "knowledge-class.md"),
        ("02-lesson-grade6.docx", ROOT / "docs" / "lesson" / "grade6.md"),
        ("04-prescription-rules.docx", ROOT / "docs" / "agents" / "prescription-rules.md"),
    ]
    for name, path in files:
        text = path.read_text(encoding="utf-8")
        first = next((ln[2:].strip() for ln in text.splitlines() if ln.startswith("# ")), path.stem)
        doc = md_to_docx(text, first)
        dest = OUT / name
        doc.save(dest)
        print("wrote", dest)

    paps = paps_doc()
    dest = OUT / "03-paps-events.docx"
    paps.save(dest)
    print("wrote", dest)


if __name__ == "__main__":
    main()
